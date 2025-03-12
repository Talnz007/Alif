from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel
# Remove the direct import of genai if not used elsewhere in *this* file
# import google.generativeai as genai
import re
import logging
from app.core.config import settings
import fitz  # PyMuPDF for PDF text extraction
import io

# --- Import the ACTUAL TextSummarizerGemini class ---
from app.endpoints.text_sumarization import TextSummarizerGemini # Correct import

# Initialize logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter()

class SummaryResponse(BaseModel):
    summary: str

text_summarizer = TextSummarizerGemini()

# New endpoint to handle file uploads for summarization
@router.post("/upload-for-summary/", response_model=SummaryResponse)
async def upload_for_summary(file: UploadFile = File(...)):
    try:
        # Read the uploaded file
        content = await file.read()

        # Extract text based on file type
        if file.content_type == "application/pdf":
            # PDF file
            text = extract_text_from_pdf(content)
        elif file.content_type == "text/plain":
            # Text file
            text = content.decode("utf-8")
        elif file.content_type.startswith("application/vnd.openxmlformats-officedocument.wordprocessingml"):
            # DOCX file - requires python-docx library
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs])
                if not text.strip(): # Check if paragraphs yielded text
                     logger.warning("DOCX contained no text in paragraphs. Trying tables (basic).")
                     # Basic table text extraction (might need improvement)
                     all_text = []
                     for table in doc.tables:
                         for row in table.rows:
                             for cell in row.cells:
                                 all_text.append(cell.text)
                     text = "\n".join(all_text)

            except ImportError:
                logger.error("python-docx library not installed. Cannot process DOCX.")
                raise HTTPException(status_code=501, detail="DOCX processing requires 'python-docx'. Please install it.")
            except Exception as docx_err:
                 logger.error(f"Error parsing DOCX file: {docx_err}")
                 raise HTTPException(status_code=500, detail="Failed to parse DOCX file.")

        else:
            logger.warning(f"Unsupported file format attempt: {file.content_type}")
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {file.content_type}")

        if not text or not text.strip():
             raise HTTPException(status_code=400, detail="Could not extract meaningful text from the file.")

        # Generate summary using existing summarizer instance
        # The `summarize_text` method comes from the imported class
        summary = await text_summarizer.summarize_text(text)
        return {"summary": summary}

    except HTTPException as http_exc:
        # Re-raise HTTPExceptions directly
        raise http_exc
    except Exception as e:
        logger.exception(f"Error during file summarization for {file.filename}: {e}") # Use logger.exception for traceback
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred during file processing: {str(e)}")

def extract_text_from_pdf(content):
    """Extract text from a PDF byte stream."""
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        text = ""
        for page_num in range(len(doc)):
            page = doc[page_num]
            text += page.get_text() + "\n" # Add newline between pages
        return text.strip() # Strip leading/trailing whitespace from final text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        # It's often better to raise the HTTPException here where the error occurs
        raise HTTPException(status_code=500, detail="Failed to extract text from PDF")
